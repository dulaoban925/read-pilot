"""DOCX文档解析器"""
from pathlib import Path
from typing import List

from docx import Document

from app.core.document_parser.base import BaseDocumentParser, ParsedDocument


class DOCXParser(BaseDocumentParser):
    """DOCX文档解析器"""

    def supports(self, file_extension: str) -> bool:
        """支持.docx文件"""
        return file_extension.lower() == ".docx"

    async def parse(self, file_path: Path) -> ParsedDocument:
        """
        解析DOCX文档

        Args:
            file_path: DOCX文件路径

        Returns:
            ParsedDocument: 解析后的文档

        Raises:
            ValueError: DOCX解析失败
        """
        try:
            doc = Document(str(file_path))

            # 提取段落
            paragraphs: List[str] = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text:
                        paragraphs.append(row_text)

            full_text = "\n\n".join(paragraphs)

            # 提取元数据
            metadata = {
                "format": "docx",
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
            }

            # 添加文档属性
            core_props = doc.core_properties
            if core_props.title:
                metadata["title"] = core_props.title
            if core_props.author:
                metadata["author"] = core_props.author
            if core_props.subject:
                metadata["subject"] = core_props.subject
            if core_props.created:
                metadata["created"] = str(core_props.created)
            if core_props.modified:
                metadata["modified"] = str(core_props.modified)

            return ParsedDocument(
                text=full_text,
                metadata=metadata,
                pages=paragraphs  # 使用段落作为"页面"
            )

        except Exception as e:
            raise ValueError(f"Failed to parse DOCX: {str(e)}")
